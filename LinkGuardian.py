"""
App Name: LinkGuardian

Description:
LinkGuardian is an application designed to facilitate online research, content writing, and link verification, ensuring that all sources used in a research process are valid and functional.

What LinkGuardian Does:
1. **Research**: Uses AI agents to investigate the latest trends in artificial intelligence (AI), logging the URLs of the sources visited during the process.
2. **Writing**: A dedicated agent writes detailed reports based on the research findings, including descriptions of the URLs used, ensuring that the mentioned content is relevant and reliable.
3. **Link Verification**: A specialized agent verifies the validity and accessibility of the URLs mentioned in the report. This involves checking the availability of the links and the relevance of the content to the AI topic.
4. **Final Report**: The final result is compiled into a Word document, divided into a main report detailing the research findings and a technical report reviewing the validity of the links.

How It Works:
- The application consists of three main agents:
  1. **Researcher**: Conducts research on AI trends and logs the accessed URLs.
  2. **Writer**: Drafts the report based on the Researcher's findings and describes the accessed URLs.
  3. **Link Revisor**: Checks the validity and accessibility of the URLs cited in the report.

- Tasks are performed sequentially, ensuring that each stage of the process is successfully completed before moving on to the next.
- The final result is saved in a Word document, which can be easily shared or used for future reference.

This application is ideal for researchers, writers, and anyone who needs to ensure the quality and validity of sources used in research papers and reports.
"""

import os
import requests
from pydantic import Field
from crewai import Agent, Task, Crew, Process
from langchain_community.llms import Ollama  # Import the Ollama module to use local models
from docx import Document  # Library to create Word documents

# 1. OpenAI API key configuration
os.environ["OPENAI_API_KEY"] = "YOUR_API_KEY_HERE"

# 2. Local model configuration with Ollama
ollama_model = Ollama(model="openhermes")

# 3. Creation of the "Researcher" agent with URL logging
class ResearcherWithURLLogging(Agent):
    visited_urls: list[str] = Field(default_factory=list)

    def perform_task(self, task):
        result = super().perform_task(task)
        # Assumes the agent accesses URLs and logs these URLs
        if hasattr(self.llm, 'visited_urls'):
            self.visited_urls.extend(self.llm.visited_urls)
        else:
            print("No URLs were captured by the LLM.")
        return result

    def get_visited_urls(self):
        return self.visited_urls

researcher = ResearcherWithURLLogging(
    role='Researcher',
    goal='Research new AI trends',
    backstory="You are an AI research assistant.",
    verbose=True,
    allow_delegation=False,
    llm=ollama_model
)

# 4. Creation of the "Writer" agent
class WriterWithCheck(Agent):
    def perform_task(self, task):
        visited_urls = task.agent.get_visited_urls()
        if visited_urls:
            # Describe URLs separately from the text
            result = super().perform_task(task)
            result += "\n\nCited URLs:\n" + "\n".join(visited_urls)
        else:
            # No URLs found, the Writer will report the problem
            result = (
                f"The agent '{task.agent.role}' could not deliver the URLs. "
                "Reason: The research did not result in URLs, or there was a failure in execution."
            )
        return result

writer = WriterWithCheck(
    role='Writer',
    goal='Write engaging blog posts about AI trends and insights.',
    backstory=(
        "You are a blog post writer specializing in AI topics and write in Portuguese. "
        "Your role includes not only writing about trends but also describing the resources and sites used in the research."
    ),
    verbose=True,
    allow_delegation=False,
    llm=ollama_model
)

# 5. New Agent: Link Revisor with Advanced Verification
class LinkRevisor(Agent):
    def check_url(self, url):
        try:
            response = requests.get(url, timeout=10)
            if response.status_code == 200:
                # Checks if the page has some expected content
                if "html" in response.headers["Content-Type"]:
                    content = response.text[:500]  # Reads the first 500 characters
                    if "AI" in content or "artificial intelligence" in content:
                        return True, "Valid, accessible link with relevant content."
                    else:
                        return False, "Accessible link, but the content may not be relevant."
                else:
                    return False, "Accessible link, but it does not appear to be a valid HTML page."
            else:
                return False, f"Invalid link. Status code: {response.status_code}"
        except requests.RequestException as e:
            return False, f"Error accessing the link: {e}"

    def perform_task(self, task):
        urls = task.agent.get_visited_urls()
        results = []
        for url in urls:
            is_valid, message = self.check_url(url)
            results.append(f"URL: {url}\nStatus: {message}\n")
        return "\n".join(results)

link_revisor = LinkRevisor(
    role='Link Revisor',
    goal='Verify the validity of the researched URLs and report issues.',
    backstory="You are responsible for ensuring that all links presented in the final report are valid and functional.",
    verbose=True,
    allow_delegation=False,
    llm=ollama_model
)

# 6. Task definitions
task1 = Task(
    description="Investigate the latest trends in AI.",
    expected_output="Detailed report on AI trends.",
    agent=researcher
)

task2 = Task(
    description=(
        "Receive the URLs researched by the Researcher and describe each site. "
        "Provide context on how these sites contributed to the research of AI trends."
    ),
    expected_output="Detailed description of the visited URLs and their relevance to the research.",
    agent=writer
)

# New Task: Link Review with Technical Report
task3 = Task(
    description="Verify the validity of the provided URLs and report any issues found.",
    expected_output="Technical review report of the links, with status of validity and accessibility.",
    agent=link_revisor
)

# 7. Team instantiation (Crew)
crew = Crew(
    agents=[researcher, writer, link_revisor],
    tasks=[task1, task2, task3],  # Tasks in sequence
    verbose=True,
    process=Process.sequential  # Executes tasks in sequence
)

# 8. Process execution with Success Condition
try:
    result = crew.kickoff()

    # 9. Saving the result to a Word document on the desktop
    document = Document()
    document.add_heading('AI Analysis Result', 0)

    # Dividing the report into two parts: Main Report and Technical Report
    main_report = str(result.tasks_output[0])  # Part 1: Main Report
    technical_report = str(result.tasks_output[1])  # Part 2: Technical Report

    document.add_paragraph(main_report)
    document.add_page_break()
    document.add_heading('Technical Review Report of the Links', level=1)
    document.add_paragraph(technical_report)

    file_path = r"D:\OneDrive\Desktop\AI_Result.docx"
    document.save(file_path)

    # 10. Printing the success message
    print(f"The result was successfully saved at: {file_path}")

except ValueError as e:
    # If any link fails, the process is interrupted, and the error is reported
    print(f"Error during execution: {e}")
