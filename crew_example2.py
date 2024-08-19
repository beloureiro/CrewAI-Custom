import os
from crewai import Agent, Task, Crew, Process
from langchain_community.llms import Ollama  # Importa o módulo Ollama para usar modelos locais
from docx import Document  # Biblioteca para criar documentos Word

# 1. Configuração da chave da API do OpenAI
os.environ["OPENAI_API_KEY"] = "SUA_CHAVE_API_AQUI"

# 2. Configuração do modelo local com Ollama
ollama_model = Ollama(model="openhermes")

# 3. Criação do agente "Researcher" (Pesquisador)
researcher = Agent(
    role='Pesquisador',
    goal='Pesquisar novas tendências em IA',
    backstory="Você é um assistente de pesquisa em IA.",
    verbose=True,
    allow_delegation=False,
    llm=ollama_model
)

# 4. Criação do agente "Writer" (Escritor)
writer = Agent(
    role='Escritor',
    goal='Escrever postagens de blog envolventes sobre tendências e insights de IA.',
    backstory="Você é um escritor de postagens de blog especializado em tópicos de IA e escreve em português.",
    verbose=True,
    allow_delegation=False,
    llm=ollama_model
)

# 5. Definição das tarefas (Tasks)
task1 = Task(
    description="Investigar as últimas tendências em IA.",
    expected_output="Relatório detalhado sobre as tendências em IA.",
    agent=researcher
)

task2 = Task(
    description="Escrever uma postagem de blog envolvente baseada nas últimas tendências em IA.",
    expected_output="Uma postagem de blog bem escrita sobre as tendências em IA.",
    agent=writer
)

# 6. Instanciação da equipe (Crew)
crew = Crew(
    agents=[researcher, writer],
    tasks=[task1, task2],
    verbose=True,
    process=Process.sequential
)

# 7. Execução do processo
result = crew.kickoff()

# 8. Salvando o resultado em um arquivo Word no desktop
document = Document()
document.add_heading('Resultado da Análise de IA', 0)

# Verifica se o resultado é uma lista e converte para string se necessário
if isinstance(result, list):
    result_str = "\n".join(result)
else:
    result_str = str(result)

document.add_paragraph(result_str)
file_path = r"D:\OneDrive\Desktop\Resultado_IA.docx"
document.save(file_path)

# 9. Impressão da mensagem de sucesso
print(f"O resultado foi salvo com sucesso em: {file_path}")
